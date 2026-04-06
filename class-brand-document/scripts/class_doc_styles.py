#!/usr/bin/env python3
"""
Class Technologies Word Document Brand Styles Module

Pre-configured python-docx style objects for creating Class-branded Word documents.
Import this module to access all brand colors, fonts, and helper functions.

Usage:
    from class_doc_styles import *

    # Create a new document
    doc = create_class_document()

    # Add elements
    add_title(doc, "My Document Title")
    add_subtitle(doc, "Subtitle | Date")
    add_section_heading(doc, "Section 1")
    add_body_text(doc, "Body paragraph text...")

    # Save
    doc.save("output.docx")
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement

# =============================================================================
# COLOR CONSTANTS
# =============================================================================

# RGBColor objects for text
NAVY = RGBColor(10, 24, 73)        # #0A1849 - Primary text, titles
PURPLE = RGBColor(71, 57, 231)     # #4739E7 - Section headings, accents
GOLD = RGBColor(255, 186, 0)       # #FFBA00 - Accent (use sparingly)
WHITE = RGBColor(255, 255, 255)    # #FFFFFF - Text on dark backgrounds

# Hex strings for XML/shading
NAVY_HEX = "0A1849"
PURPLE_HEX = "4739E7"
GOLD_HEX = "FFBA00"
WHITE_HEX = "FFFFFF"
LIGHT_PURPLE_HEX = "EBE9FC"        # Table header alt style
LIGHT_GRAY_HEX = "F0F0F5"          # Zebra stripe rows
FAQ_BG_HEX = "F8F7FE"              # FAQ question cells
BORDER_GRAY_HEX = "E0E0E0"         # Table borders

# =============================================================================
# TYPOGRAPHY CONSTANTS
# =============================================================================

# Font names
FONT_PRIMARY = "Roboto"
FONT_BODY = "Roboto Light"

# Font sizes
SIZE_TITLE = Pt(24)
SIZE_SUBTITLE = Pt(18)
SIZE_H1 = Pt(14)
SIZE_H2 = Pt(12)
SIZE_H3 = Pt(10)
SIZE_BODY = Pt(10)
SIZE_BULLET = Pt(9)
SIZE_TABLE = Pt(10)

# =============================================================================
# SPACING CONSTANTS
# =============================================================================

# Page margins
MARGIN_TOP = Inches(0.5)
MARGIN_BOTTOM = Inches(0.75)
MARGIN_LEFT = Inches(0.75)
MARGIN_RIGHT = Inches(0.75)

# Paragraph spacing
LINE_SPACING = 1.15
SPACE_TITLE_AFTER = Pt(6)
SPACE_SUBTITLE_AFTER = Pt(18)
SPACE_H1_BEFORE = Pt(18)
SPACE_H1_AFTER = Pt(12)
SPACE_H2_BEFORE = Pt(12)
SPACE_H2_AFTER = Pt(6)
SPACE_H3_BEFORE = Pt(10)
SPACE_H3_AFTER = Pt(4)
SPACE_BODY_AFTER = Pt(6)

# =============================================================================
# DOCUMENT CREATION
# =============================================================================

def create_class_document():
    """Create a new document with Class brand settings."""
    doc = Document()

    section = doc.sections[0]
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT

    return doc


# =============================================================================
# HEADING FUNCTIONS
# =============================================================================

def add_title(doc, text):
    """Add document title (24pt, Bold, Navy, Centered)."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = FONT_PRIMARY
    run.font.size = SIZE_TITLE
    run.font.bold = True
    run.font.color.rgb = NAVY
    para.paragraph_format.space_after = SPACE_TITLE_AFTER
    return para


def add_subtitle(doc, text):
    """Add subtitle (18pt, Regular, Purple, Centered)."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = FONT_PRIMARY
    run.font.size = SIZE_SUBTITLE
    run.font.bold = False
    run.font.color.rgb = PURPLE
    para.paragraph_format.space_after = SPACE_SUBTITLE_AFTER
    return para


def add_section_heading(doc, text):
    """Add H1 section heading (14pt, Bold, Purple)."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.font.name = FONT_PRIMARY
    run.font.size = SIZE_H1
    run.font.bold = True
    run.font.color.rgb = PURPLE
    para.paragraph_format.space_before = SPACE_H1_BEFORE
    para.paragraph_format.space_after = SPACE_H1_AFTER
    return para


def add_subsection_heading(doc, text):
    """Add H2 subsection heading (12pt, Bold, Navy)."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.font.name = FONT_PRIMARY
    run.font.size = SIZE_H2
    run.font.bold = True
    run.font.color.rgb = NAVY
    para.paragraph_format.space_before = SPACE_H2_BEFORE
    para.paragraph_format.space_after = SPACE_H2_AFTER
    return para


def add_h3_heading(doc, text):
    """Add H3 heading (10pt, Bold, Purple)."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.font.name = FONT_PRIMARY
    run.font.size = SIZE_H3
    run.font.bold = True
    run.font.color.rgb = PURPLE
    para.paragraph_format.space_before = SPACE_H3_BEFORE
    para.paragraph_format.space_after = SPACE_H3_AFTER
    return para


# =============================================================================
# BODY TEXT FUNCTIONS
# =============================================================================

def add_body_text(doc, text):
    """Add body paragraph (10pt, Roboto Light, Navy)."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = SIZE_BODY
    run.font.bold = False
    run.font.color.rgb = NAVY
    para.paragraph_format.line_spacing = LINE_SPACING
    para.paragraph_format.space_after = SPACE_BODY_AFTER
    return para


def add_bullet_point(doc, text, level=0):
    """Add bullet point (9pt, Roboto Light, Navy with Purple bullet)."""
    para = doc.add_paragraph()

    # Purple bullet
    bullet_run = para.add_run("•  ")
    bullet_run.font.name = FONT_PRIMARY
    bullet_run.font.size = SIZE_BULLET
    bullet_run.font.color.rgb = PURPLE

    # Navy text
    text_run = para.add_run(text)
    text_run.font.name = FONT_BODY
    text_run.font.size = SIZE_BULLET
    text_run.font.color.rgb = NAVY

    para.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return para


def add_numbered_item(doc, number, text):
    """Add numbered item (Purple number, Navy text)."""
    para = doc.add_paragraph()

    # Purple number
    num_run = para.add_run(f"{number}. ")
    num_run.font.name = FONT_PRIMARY
    num_run.font.size = SIZE_BODY
    num_run.font.bold = True
    num_run.font.color.rgb = PURPLE

    # Navy text
    text_run = para.add_run(text)
    text_run.font.name = FONT_BODY
    text_run.font.size = SIZE_BODY
    text_run.font.color.rgb = NAVY

    return para


# =============================================================================
# TABLE FUNCTIONS
# =============================================================================

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def style_cell_text(cell, font_name, size, bold, color):
    """Style text in a table cell."""
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = font_name
            run.font.size = size
            run.font.bold = bold
            run.font.color.rgb = color


def add_purple_header_table(doc, headers, data):
    """Create table with purple header row and zebra striped data."""
    table = doc.add_table(rows=1 + len(data), cols=len(headers))

    # Header row
    header_row = table.rows[0]
    for idx, header_text in enumerate(headers):
        cell = header_row.cells[idx]
        cell.text = header_text
        set_cell_shading(cell, PURPLE_HEX)
        style_cell_text(cell, FONT_PRIMARY, SIZE_TABLE, True, WHITE)

    # Data rows with zebra striping
    for row_idx, row_data in enumerate(data):
        row = table.rows[row_idx + 1]
        fill = WHITE_HEX if row_idx % 2 == 0 else LIGHT_GRAY_HEX

        for col_idx, cell_value in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(cell_value)
            set_cell_shading(cell, fill)
            style_cell_text(cell, FONT_BODY, SIZE_TABLE, False, NAVY)

    return table


def add_light_purple_header_table(doc, headers, data):
    """Create table with light purple header row."""
    table = doc.add_table(rows=1 + len(data), cols=len(headers))

    # Header row
    header_row = table.rows[0]
    for idx, header_text in enumerate(headers):
        cell = header_row.cells[idx]
        cell.text = header_text
        set_cell_shading(cell, LIGHT_PURPLE_HEX)
        style_cell_text(cell, FONT_PRIMARY, SIZE_TABLE, True, NAVY)

    # Data rows
    for row_idx, row_data in enumerate(data):
        row = table.rows[row_idx + 1]

        for col_idx, cell_value in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(cell_value)
            style_cell_text(cell, FONT_BODY, SIZE_TABLE, False, NAVY)

    return table


def add_faq_table(doc, qa_pairs):
    """Create FAQ-style table with question/answer columns.

    Args:
        qa_pairs: List of tuples [(question, answer), ...]
    """
    table = doc.add_table(rows=len(qa_pairs), cols=2)

    for row_idx, (question, answer) in enumerate(qa_pairs):
        row = table.rows[row_idx]

        # Question cell (light purple bg, purple text)
        q_cell = row.cells[0]
        q_cell.text = question
        set_cell_shading(q_cell, FAQ_BG_HEX)
        for para in q_cell.paragraphs:
            for run in para.runs:
                run.font.name = FONT_PRIMARY
                run.font.size = SIZE_TABLE
                run.font.color.rgb = PURPLE

        # Answer cell (white bg, navy text)
        a_cell = row.cells[1]
        a_cell.text = answer
        for para in a_cell.paragraphs:
            for run in para.runs:
                run.font.name = FONT_BODY
                run.font.size = SIZE_TABLE
                run.font.color.rgb = NAVY

    return table


def set_table_borders(table, color=BORDER_GRAY_HEX, size="4"):
    """Add thin borders to all table cells."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:color'), color)
        tblBorders.append(border)

    tblPr.append(tblBorders)


# =============================================================================
# COMPLETE DOCUMENT TEMPLATES
# =============================================================================

def create_faq_document(title, subtitle, sections):
    """Create a complete FAQ document.

    Args:
        title: Document title
        subtitle: Subtitle with date/context
        sections: Dict of {section_name: [(q, a), ...]}

    Returns:
        Document object
    """
    doc = create_class_document()

    add_title(doc, title)
    add_subtitle(doc, subtitle)

    # Table of Contents
    if len(sections) > 1:
        add_section_heading(doc, "Table of Contents")
        for idx, section_name in enumerate(sections.keys(), 1):
            add_body_text(doc, f"{idx}. {section_name}")
        doc.add_paragraph()

    # Sections with Q&A
    for section_name, qa_pairs in sections.items():
        add_section_heading(doc, section_name)
        table = add_faq_table(doc, qa_pairs)
        set_table_borders(table)
        doc.add_paragraph()

    return doc


def create_roadmap_document(title, subtitle, sections):
    """Create a product roadmap document.

    Args:
        title: Document title
        subtitle: Subtitle
        sections: Dict of {section_name: {subsection: {feature_group: [features]}}}

    Returns:
        Document object
    """
    doc = create_class_document()

    add_title(doc, title)
    add_subtitle(doc, subtitle)

    for section_name, subsections in sections.items():
        add_section_heading(doc, section_name.upper())

        for subsection_name, feature_groups in subsections.items():
            add_subsection_heading(doc, subsection_name)

            for feature_group, features in feature_groups.items():
                add_h3_heading(doc, feature_group)

                for feature in features:
                    add_bullet_point(doc, feature)

    return doc


def create_report_document(title, date, sections):
    """Create a standard report document.

    Args:
        title: Report title
        date: Date/context string
        sections: List of section dicts with keys:
            - heading: Section title
            - content: List of paragraphs (optional)
            - bullets: List of bullet points (optional)
            - subsections: List of subsection dicts (optional)
            - table: Dict with 'headers' and 'data' (optional)

    Returns:
        Document object
    """
    doc = create_class_document()

    add_title(doc, title)
    add_subtitle(doc, date)

    for sect in sections:
        add_section_heading(doc, sect['heading'])

        if 'content' in sect:
            for para_text in sect['content']:
                add_body_text(doc, para_text)

        if 'subsections' in sect:
            for subsect in sect['subsections']:
                add_subsection_heading(doc, subsect['heading'])
                if 'content' in subsect:
                    add_body_text(doc, subsect['content'])

        if 'bullets' in sect:
            for bullet_text in sect['bullets']:
                add_bullet_point(doc, bullet_text)

        if 'table' in sect:
            doc.add_paragraph()
            table = add_purple_header_table(
                doc,
                sect['table']['headers'],
                sect['table']['data']
            )
            set_table_borders(table)

    return doc


# =============================================================================
# MODULE TEST
# =============================================================================

if __name__ == "__main__":
    print("Class Document Styles Module")
    print("=" * 40)
    print(f"Navy: #{NAVY_HEX}")
    print(f"Purple: #{PURPLE_HEX}")
    print(f"Gold: #{GOLD_HEX}")
    print(f"\nPrimary Font: {FONT_PRIMARY}")
    print(f"Body Font: {FONT_BODY}")
    print("\nImport this module to use Class brand styles in python-docx.")
